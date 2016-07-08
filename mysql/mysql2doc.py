#!/usr/bin/env python
# -*- coding: utf-8 -*-

import sys
import datetime

from docx import Document

from mysql import mysql_select, mysql_init

@mysql_select('SELECT TABLE_NAME, ENGINE, TABLE_COLLATION, TABLE_COMMENT FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_SCHEMA=%s ORDER BY TABLE_NAME')
def table_schema(schema):
  pass

@mysql_select('SELECT TABLE_NAME,COLUMN_NAME,COLUMN_TYPE,COLUMN_COMMENT,COLUMN_KEY FROM INFORMATION_SCHEMA.COLUMNS WHERE TABLE_SCHEMA=%s ORDER BY TABLE_NAME')
def table_colume(schema):
  pass
  
def doit():

  host = '192.168.1.254'
  user = 'root'
  password = 'renrenbx'
  database = 'renrenbx'
  
  mysql_init(mysql_host = host, mysql_database = database, mysql_user = user, mysql_password = password)  
  
  tables = {}
  
  for column in table_schema(database):
    tables[column['TABLE_NAME']] = {'info':column,'columns':[]} 
        
  for column in table_colume(database):
    tables[column['TABLE_NAME']]['columns'] += [column]  
    
  document = Document()
  document.add_heading(database, 0)
  
  i = 0
  max = len(tables)
  
  for key in sorted(tables.keys()):
    
    i = i + 1
    value = int(round((i * 1.0) / max * 100))
    
    sys.stdout.write(' [' + '#' * i + '] %s%%' % value + '\r')
    sys.stdout.flush()
    
    document.add_heading(key, 1)
    table_engine = tables[key]['info']['ENGINE']
    paragraph = document.add_paragraph()
    paragraph.add_run(table_engine).bold = True
    table_comment = tables[key]['info']['TABLE_COMMENT']
    paragraph = document.add_paragraph()
    paragraph.add_run(table_comment if table_comment else u'无注释').bold = True
    table = document.add_table(rows = 1, cols = 4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'字段'
    hdr_cells[1].text = u'主键'
    hdr_cells[2].text = u'类型'
    hdr_cells[3].text = u'注释'
    for column in tables[key]['columns']:
      row_cell = table.add_row().cells
      row_cell[0].text = column['COLUMN_NAME'] 
      row_cell[1].text = column['COLUMN_KEY'] if column['COLUMN_KEY'] else '-'
      row_cell[2].text = column['COLUMN_TYPE'] 
      row_cell[3].text = column['COLUMN_COMMENT'] if column['COLUMN_COMMENT'] else '-'
  
    document.save('%s-%s.docx' % (database,datetime.datetime.now().strftime("%Y%m%d%H")))
    
    
if __name__ == '__main__':
  doit()